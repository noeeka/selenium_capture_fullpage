import codecs
import datetime
import smtplib
import sys
import time
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from docx import Document
from docx.shared import Inches
from selenium import webdriver
from win32com.client import Dispatch

sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())
starttime = datetime.datetime.now()

from_addr = ''
password = ''
to_addrs = ['', '']

lists = ['bh', 'bd', 'bn', 'in', 'id', 'my', 'pk', 'sg', 'lk', 'vn', 'hk', 'cn', 'tw', 'ci', 'fk', 'gm', 'gh', 'je',
         'je', 'jo', 'ke', 'np', 'ng', 'sl', 'tz', 'ae', 'ug', 'zm', 'bw', 'zw']

document = Document()
for x in lists:
    url = "https://www.sc.com/" + x

    image_desktop_full = "image_desktop_full.png"
    image_mobile_full = "image_mobile_full.png"
    option = webdriver.ChromeOptions()

    # capture for desktop option
    option.add_argument('--headless')
    option.add_argument('--disable-gpu')
    option.add_argument("--window-size=1920,1080")
    option.add_argument("--hide-scrollbars")

    driver = webdriver.Chrome(chrome_options=option)

    driver.get(url)
    print(driver.title)
    print(url)
    time.sleep(8)
    scroll_width = driver.execute_script('return document.body.parentNode.scrollWidth')
    scroll_height = driver.execute_script('return document.body.parentNode.scrollHeight')
    driver.set_window_size(scroll_width, scroll_height)
    driver.save_screenshot(image_desktop_full)

    # capture for mobile option
    option.add_argument('--headless')
    option.add_argument('--disable-gpu')
    option.add_argument("--window-size=390,884")
    option.add_argument("--hide-scrollbars")

    driver = webdriver.Chrome(chrome_options=option)

    driver.get(url)
    print(driver.title)
    print(url)
    time.sleep(8)
    scroll_width = driver.execute_script('return document.body.parentNode.scrollWidth')
    scroll_height = driver.execute_script('return document.body.parentNode.scrollHeight')
    driver.set_window_size(scroll_width, scroll_height)
    driver.save_screenshot(image_mobile_full)
    driver.quit()

    # 初始化建立第一个自然段
    pl = document.add_paragraph()
    # 对齐方式为居中，没有这句话默认左对齐
    run1 = pl.add_run(url)

    document.sections[0].left_margin = Inches(0.3)
    document.sections[0].bottom_margin = Inches(0.3)

    document.add_picture(image_desktop_full, height=Inches(9))

    new_section = document.add_section()
    document.add_picture(image_mobile_full, height=Inches(9))

# document.add_page_break()
document.save('report.docx')

word = Dispatch('Word.Application')
doc = word.Documents.Open('D:\\Project\\selenium_capture_fullpage\\report.docx')
doc.SaveAs('D:\\Project\\selenium_capture_fullpage\\report.pdf', 17)
doc.Close()
word.Quit()



content = 'hello, PFA.'
text_apart = MIMEText(content)

pdf_file = 'report.pdf'
pdf_apart = MIMEApplication(open(pdf_file, 'rb').read())
pdf_apart.add_header('Content-Disposition', 'attachment', filename=pdf_file)

m = MIMEMultipart()
m.attach(text_apart)
m.attach(pdf_apart)
m['Subject'] = 'all_market_homepage_fullpage'

try:
    server = smtplib.SMTP('smtp.126.com')
    server.login(from_addr, password)
    server.sendmail(from_addr, to_addrs, m.as_string())
    print('send success')
    server.quit()
except smtplib.SMTPException as e:
    print('error:', e)  # 打印错误
endtime = datetime.datetime.now()
print("In total:", len(lists))
print("Total running time is :", ((endtime - starttime).seconds) / 60, "mins")
